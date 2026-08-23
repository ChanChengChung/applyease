from io import BytesIO
from uuid import uuid4

from docx import Document
from fastapi.testclient import TestClient
from pypdf import PdfReader

from app.main import app


def _register(client: TestClient, email: str) -> str:
    response = client.post(
        "/api/v1/auth/register",
        headers={"X-ApplyEase-Client": "browser-extension"},
        json={"email": email, "password": "export-secure-password"},
    )
    assert response.status_code == 201, response.text

    return response.json()["access_token"]


def _resume(client: TestClient, headers: dict[str, str]):
    suffix = uuid4().hex[:8]

    experience = client.post(
        "/api/v1/experiences",
        headers=headers,
        json={
            "title": f"AI Developer {suffix}",
            "organization": f"Student Lab {suffix}",
            "description": "Built a Python and React application using confirmed project evidence.",
            "skills": ["Python", "React"],
            "achievements": [],
            "source_file": "manual",
            "confirmed": True,
        },
    )

    assert experience.status_code == 200, experience.text

    job = client.post(
        "/api/v1/jobs/analyze",
        headers=headers,
        json={
            "title": f"AI Engineering Intern {suffix}",
            "company": "Polymer Capital",
            "description": "Build Python and React tools with an engineering team.",
        },
    )

    assert job.status_code == 200, job.text

    resume = client.post(
        f"/api/v1/materials/resume/generate?job_id={job.json()['id']}", headers=headers
    )

    assert resume.status_code == 200 and resume.json()["fact_check_passed"] is True

    return resume.json(), job.json()


def test_all_resume_templates_export_valid_docx_and_pdf_with_safe_headers():
    client = TestClient(app)
    token = _register(client, f"export-{uuid4()}@example.com")

    headers = {"Authorization": f"Bearer {token}"}
    resume, job = _resume(client, headers)

    for template in ("classic", "modern", "compact"):
        docx_response = client.post(
            f"/api/v1/materials/{resume['id']}/export",
            headers=headers,
            json={
                "format": "docx",
                "template": template,
                "include_sources": True,
                "display_name": "Chen Zhengzhong",
                "contact_line": "chen@example.com · github.com/chen",
                "email": "chen@example.com",
                "phone": "+852 6406 1561",
                "location": "Hong Kong SAR",
                "linkedin_url": "linkedin.com/in/chen",
                "github_url": "github.com/chen",
                "font_style": "serif",
                "density": "relaxed",
                "accent": "navy",
            },
        )

        assert docx_response.status_code == 200 and docx_response.content.startswith(b"PK")

        assert docx_response.headers["content-type"].startswith("application/vnd.openxmlformats")

        assert f"-{template}.docx" in docx_response.headers["content-disposition"]

        document = Document(BytesIO(docx_response.content))

        text = "\n".join(item.text for item in document.paragraphs)

        assert "Chen Zhengzhong" in text and "chen@example.com" in text
        assert "+852 6406 1561" in text and "LinkedIn: linkedin.com/in/chen" in text
        assert document.paragraphs[0].runs[0].font.name == "Times New Roman"

        assert job["title"] in text and "ApplyEase Evidence Appendix" in text

        section = document.sections[0]

        assert (
            round(section.page_width.inches, 1) == 8.5
            and round(section.page_height.inches, 1) == 11.0
        )

        pdf_response = client.post(
            f"/api/v1/materials/{resume['id']}/export",
            headers=headers,
            json={
                "format": "pdf",
                "template": template,
                "display_name": "Chen Zhengzhong",
                "font_style": "sans",
                "density": "compact",
                "accent": "black",
            },
        )

        assert pdf_response.status_code == 200 and pdf_response.content.startswith(b"%PDF")

        assert pdf_response.headers["content-type"].startswith("application/pdf")

        reader = PdfReader(BytesIO(pdf_response.content))
        assert len(reader.pages) >= 1

        assert (
            round(float(reader.pages[0].mediabox.width)) == 612
            and round(float(reader.pages[0].mediabox.height)) == 792
        )


def test_export_rejects_wrong_type_invalid_options_and_failed_fact_check():
    client = TestClient(app)
    token = _register(client, f"export-boundary-{uuid4()}@example.com")

    headers = {"Authorization": f"Bearer {token}"}
    resume, job = _resume(client, headers)

    assert (
        client.post(
            f"/api/v1/materials/{resume['id']}/export",
            headers=headers,
            json={"format": "zip", "template": "classic", "display_name": "Test User"},
        ).status_code
        == 422
    )

    assert (
        client.post(
            f"/api/v1/materials/{resume['id']}/export",
            headers=headers,
            json={
                "format": "pdf",
                "template": "classic",
                "display_name": "Test User",
                "font_style": "comic-sans",
            },
        ).status_code
        == 422
    )

    assert (
        client.post(
            f"/api/v1/materials/{resume['id']}/export",
            headers=headers,
            json={"format": "pdf", "template": "unknown", "display_name": "Test User"},
        ).status_code
        == 422
    )

    assert (
        client.post(
            f"/api/v1/materials/{resume['id']}/export",
            headers=headers,
            json={"format": "pdf", "template": "classic", "display_name": "   "},
        ).status_code
        == 422
    )

    edited = client.patch(
        f"/api/v1/materials/{resume['id']}",
        headers=headers,
        json={"text": resume["text"] + "\nImproved results by 999%."},
    )

    assert edited.status_code == 200 and edited.json()["fact_check_passed"] is False

    blocked = client.post(
        f"/api/v1/materials/{resume['id']}/export",
        headers=headers,
        json={"format": "pdf", "template": "modern", "display_name": "Test User"},
    )

    assert blocked.status_code == 409

    cover = client.post(
        f"/api/v1/materials/cover-letter/generate?job_id={job['id']}", headers=headers
    )

    wrong_type = client.post(
        f"/api/v1/materials/{cover.json()['id']}/export",
        headers=headers,
        json={"format": "docx", "template": "classic", "display_name": "Test User"},
    )

    assert wrong_type.status_code == 400


def test_resume_export_respects_user_ownership():
    first = TestClient(app)
    second = TestClient(app)

    first_token = _register(first, f"export-owner-a-{uuid4()}@example.com")

    second_token = _register(second, f"export-owner-b-{uuid4()}@example.com")

    resume, _ = _resume(first, {"Authorization": f"Bearer {first_token}"})

    hidden = second.post(
        f"/api/v1/materials/{resume['id']}/export",
        headers={"Authorization": f"Bearer {second_token}"},
        json={"format": "pdf", "template": "classic", "display_name": "Test User"},
    )
    assert hidden.status_code == 404
