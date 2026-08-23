from __future__ import annotations

import re
from html import escape
from io import BytesIO
from typing import Literal

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import HRFlowable, PageBreak, Paragraph, SimpleDocTemplate, Spacer

from app.models.job import Job
from app.models.material import GeneratedMaterial

ResumeTemplate = Literal["classic", "modern", "compact"]
ResumeFontStyle = Literal["default", "sans", "serif", "microsoft_yahei"]
ResumeDensity = Literal["relaxed", "standard", "compact"]
ResumeAccent = Literal["template", "navy", "black"]

TEMPLATES = {
    "classic": {
        "label": "Classic",
        "font": "Times New Roman",
        "pdf_font": "Times-Roman",
        "accent": "222222",
        "title": 20,
        "subtitle": 10,
        "body": 10.5,
        "heading": 11.5,
        "margin": 0.72,
        "after": 5,
        "line": 1.08,
    },
    "modern": {
        "label": "Modern",
        "font": "Arial",
        "pdf_font": "Helvetica",
        "accent": "173F5F",
        "title": 22,
        "subtitle": 10,
        "body": 10.3,
        "heading": 12,
        "margin": 0.68,
        "after": 5,
        "line": 1.1,
    },
    "compact": {
        "label": "Compact",
        "font": "Arial",
        "pdf_font": "Helvetica",
        "accent": "1F4D78",
        "title": 18,
        "subtitle": 9,
        "body": 9.2,
        "heading": 10.5,
        "margin": 0.55,
        "after": 3,
        "line": 1.0,
    },
}


def _export_config(
    template: ResumeTemplate,
    font_style: ResumeFontStyle = "default",
    density: ResumeDensity = "standard",
    accent: ResumeAccent = "template",
) -> dict:
    """Return a presentation-only export configuration.

    These options are deliberately resolved at export time instead of being
    persisted into generated material: a student can experiment freely
    without changing a fact-grounded resume version.
    """
    config = dict(TEMPLATES[template])
    if font_style == "sans":
        config.update(font="Arial", pdf_font="Helvetica")
    elif font_style == "serif":
        config.update(font="Times New Roman", pdf_font="Times-Roman")
    elif font_style == "microsoft_yahei":
        # DOCX will use the requested system font. ReportLab uses an embedded
        # CJK fallback for Chinese runs (see _pdf_markup) so PDF downloads
        # remain portable even where Microsoft YaHei is not installed.
        config.update(font="Microsoft YaHei", pdf_font="Helvetica")

    if accent == "navy":
        config["accent"] = "173F5F"
    elif accent == "black":
        config["accent"] = "222222"

    if density == "relaxed":
        config["margin"] = min(float(config["margin"]) + 0.08, 0.85)
        config["after"] = float(config["after"]) + 1.5
        config["line"] = float(config["line"]) + 0.12
    elif density == "compact":
        config["margin"] = max(float(config["margin"]) - 0.08, 0.45)
        config["body"] = max(float(config["body"]) - 0.7, 8.5)
        config["after"] = max(float(config["after"]) - 1.5, 1.5)
        config["line"] = max(float(config["line"]) - 0.08, 0.96)
    return config


def _safe_filename(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "-", value.strip()).strip("-.")

    return cleaned[:80] or "resume"


def export_filename(job: Job, template: ResumeTemplate, extension: str) -> str:
    role = _safe_filename(job.title)

    company = _safe_filename(job.company) if job.company else "application"

    return f"ApplyEase-{company}-{role}-{template}.{extension}"


def _resume_lines(record: GeneratedMaterial) -> list[str]:
    text = str((record.content or {}).get("text", "")).replace("\r\n", "\n").strip()

    if not text:

        raise ValueError("Resume material is empty")

    return [line.strip() for line in text.splitlines()]


def resume_sections(record: GeneratedMaterial) -> list[tuple[str, list[str]]]:
    """Split material text into stable, exportable heading sections without inventing content."""

    sections: list[tuple[str, list[str]]] = []

    preface: list[str] = []

    current_name: str | None = None

    current_lines: list[str] = []

    for line in _resume_lines(record):

        if _is_heading(line):

            if current_name is not None:
                sections.append((current_name, current_lines))
            current_name, current_lines = line.rstrip(":"), [line]

        elif current_name is None:
            preface.append(line)

        else:
            current_lines.append(line)

    if current_name is not None:
        sections.append((current_name, current_lines))

    if preface:
        sections.insert(0, ("Resume summary", preface))

    return sections or [("Resume", _resume_lines(record))]


def resume_lines_for_export(
    record: GeneratedMaterial,
    section_order: list[str] | None = None,
    hidden_sections: list[str] | None = None,
) -> list[str]:
    sections = resume_sections(record)

    by_name = {name: lines for name, lines in sections}

    hidden = set(hidden_sections or [])

    ordered = [name for name in (section_order or []) if name in by_name]

    ordered.extend(name for name, _ in sections if name not in ordered)

    lines = [line for name in ordered if name not in hidden for line in by_name[name]]

    if not lines:

        raise ValueError("Select at least one resume section before exporting")

    return lines


def _is_heading(line: str) -> bool:
    compact = line.strip().rstrip(":")

    return bool(
        compact
        and len(compact) <= 50
        and compact.upper() == compact
        and any(char.isalpha() for char in compact)
    )


def _set_run_font(
    run, name: str, size: float, color: str = "222222", bold: bool = False, italic: bool = False
) -> None:
    run.font.name = name

    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)

    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)

    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)

    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)

    run.bold = bold
    run.italic = italic


def _bottom_border(paragraph, color: str, size: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))

    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    border = OxmlElement("w:bottom")

    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), "2")
    border.set(qn("w:color"), color)

    borders.append(border)


def _configure_docx(doc: Document, config: dict) -> None:
    section = doc.sections[0]

    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(config["margin"]))
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]

    normal.font.name = config["font"]
    normal.font.size = Pt(config["body"])

    normal._element.rPr.rFonts.set(qn("w:ascii"), config["font"])
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), config["font"])

    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(config["after"])
    normal.paragraph_format.line_spacing = config["line"]


def _contact_rows(
    contact_line: str,
    email: str,
    phone: str,
    location: str,
    linkedin_url: str,
    github_url: str,
) -> list[str]:
    """Build compact, labelled contact rows from deliberate user fields.

    ``contact_line`` remains as a legacy fallback for old saved profiles, but
    a structured profile takes priority so the exported resume is legible.
    """
    primary = [
        value
        for value in (
            location,
            f"Phone: {phone}" if phone else "",
            f"Email: {email}" if email else "",
        )
        if value
    ]
    social = [
        f"LinkedIn: {linkedin_url}" if linkedin_url else "",
        f"GitHub: {github_url}" if github_url else "",
    ]
    rows = ["  |  ".join(primary)] if primary else []
    social_row = "  |  ".join(item for item in social if item)
    if social_row:
        rows.append(social_row)
    if not rows and contact_line:
        rows.append(contact_line)
    return rows


def build_resume_docx(
    record: GeneratedMaterial,
    job: Job,
    template: ResumeTemplate,
    include_sources: bool = False,
    display_name: str = "[YOUR NAME]",
    contact_line: str = "",
    email: str = "",
    phone: str = "",
    location: str = "",
    linkedin_url: str = "",
    github_url: str = "",
    section_order: list[str] | None = None,
    hidden_sections: list[str] | None = None,
    font_style: ResumeFontStyle = "default",
    density: ResumeDensity = "standard",
    accent: ResumeAccent = "template",
) -> bytes:
    config = _export_config(template, font_style, density, accent)
    doc = Document()
    _configure_docx(doc, config)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title.paragraph_format.space_after = Pt(2)

    _set_run_font(
        title.add_run(display_name), config["font"], config["title"], config["accent"], bold=True
    )

    for row in _contact_rows(contact_line, email, phone, location, linkedin_url, github_url):
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.paragraph_format.space_after = Pt(1)

        _set_run_font(contact.add_run(row), config["font"], config["subtitle"], "444444")
    divider = doc.add_paragraph()
    divider.paragraph_format.space_after = Pt(7)

    if template == "modern":
        _bottom_border(divider, config["accent"], 14)

    for line in resume_lines_for_export(record, section_order, hidden_sections):

        if not line:
            continue

        paragraph = doc.add_paragraph()

        if _is_heading(line):
            paragraph.paragraph_format.space_before = Pt(7 if template != "compact" else 4)

            paragraph.paragraph_format.space_after = Pt(3)

            _set_run_font(
                paragraph.add_run(line.rstrip(":")),
                config["font"],
                config["heading"],
                config["accent"],
                bold=True,
            )

            _bottom_border(paragraph, "B8C4CE" if template != "classic" else "777777", 5)

        elif line.startswith(("- ", "• ")):
            paragraph.style = doc.styles["List Bullet"]

            paragraph.paragraph_format.left_indent = Inches(0.22)
            paragraph.paragraph_format.first_line_indent = Inches(-0.15)

            _set_run_font(paragraph.add_run(line[2:].strip()), config["font"], config["body"])

        else:
            _set_run_font(
                paragraph.add_run(line),
                config["font"],
                config["body"],
                bold=" | " in line and not line.lower().startswith("skills:"),
            )

    if include_sources:
        doc.add_section(WD_SECTION_START.NEW_PAGE)

        heading = doc.add_paragraph()
        _set_run_font(
            heading.add_run("ApplyEase Evidence Appendix"),
            config["font"],
            16,
            config["accent"],
            bold=True,
        )

        note = doc.add_paragraph()
        _set_run_font(
            note.add_run(
                "This appendix is for review and provenance; remove it before submitting the resume."
            ),
            config["font"],
            9,
            "666666",
            italic=True,
        )

        for source in (record.content or {}).get("sources", []):

            if not isinstance(source, dict):
                continue

            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.16)

            label = f"Experience #{source.get('experience_id')}: {source.get('experience_title', 'Confirmed experience')}"

            _set_run_font(p.add_run(label), config["font"], 9.5, config["accent"], bold=True)

            detail = doc.add_paragraph()
            detail.paragraph_format.left_indent = Inches(0.25)

            _set_run_font(detail.add_run(str(source.get("text", ""))), config["font"], 9)

    properties = doc.core_properties
    properties.title = f"Resume - {job.title}"
    properties.subject = "ApplyEase evidence-grounded resume export"
    properties.author = "ApplyEase"

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def _pdf_styles(config: dict):

    try:
        pdfmetrics.getFont(config["pdf_font"])

    except KeyError:
        pdfmetrics.registerFont(UnicodeCIDFont(config["pdf_font"]))

    try:
        pdfmetrics.getFont("MSung-Light")

    except KeyError:
        pdfmetrics.registerFont(UnicodeCIDFont("MSung-Light"))
    font = config["pdf_font"]
    accent = colors.HexColor("#" + config["accent"])

    return {
        "title": ParagraphStyle(
            "ResumeTitle",
            fontName=font,
            fontSize=config["title"],
            leading=config["title"] + 2,
            textColor=accent,
            alignment=TA_CENTER,
            spaceAfter=3,
        ),
        "meta": ParagraphStyle(
            "ResumeMeta",
            fontName=font,
            fontSize=config["subtitle"],
            leading=config["subtitle"] + 2,
            textColor=colors.HexColor("#555555"),
            alignment=TA_CENTER,
            spaceAfter=8,
        ),
        "heading": ParagraphStyle(
            "ResumeHeading",
            fontName=font,
            fontSize=config["heading"],
            leading=config["heading"] + 2,
            textColor=accent,
            spaceBefore=7,
            spaceAfter=3,
        ),
        "body": ParagraphStyle(
            "ResumeBody",
            fontName=font,
            fontSize=config["body"],
            leading=config["body"] * config["line"] + 2,
            textColor=colors.HexColor("#222222"),
            spaceAfter=config["after"],
        ),
        "source": ParagraphStyle(
            "ResumeSource",
            fontName=font,
            fontSize=9,
            leading=11,
            textColor=colors.HexColor("#333333"),
            leftIndent=12,
            firstLineIndent=-8,
            spaceAfter=5,
        ),
    }


def _pdf_markup(value: str) -> str:
    """Escape Paragraph markup and use a Traditional Chinese fallback only for CJK runs."""

    parts = re.split(r"([\u3400-\u9fff\uf900-\ufaff]+)", value)

    return "".join(
        (
            f'<font name="MSung-Light">{escape(part)}</font>'
            if re.search(r"[\u3400-\u9fff\uf900-\ufaff]", part)
            else escape(part)
        )
        for part in parts
    )


def build_resume_pdf(
    record: GeneratedMaterial,
    job: Job,
    template: ResumeTemplate,
    include_sources: bool = False,
    display_name: str = "[YOUR NAME]",
    contact_line: str = "",
    email: str = "",
    phone: str = "",
    location: str = "",
    linkedin_url: str = "",
    github_url: str = "",
    section_order: list[str] | None = None,
    hidden_sections: list[str] | None = None,
    font_style: ResumeFontStyle = "default",
    density: ResumeDensity = "standard",
    accent: ResumeAccent = "template",
) -> bytes:
    config = _export_config(template, font_style, density, accent)
    styles = _pdf_styles(config)
    output = BytesIO()

    doc = SimpleDocTemplate(
        output,
        pagesize=letter,
        leftMargin=config["margin"] * inch,
        rightMargin=config["margin"] * inch,
        topMargin=config["margin"] * inch,
        bottomMargin=config["margin"] * inch,
        title=f"Resume - {job.title}",
        author="ApplyEase",
    )
    story = [Paragraph(_pdf_markup(display_name), styles["title"])]

    for row in _contact_rows(contact_line, email, phone, location, linkedin_url, github_url):
        story.append(Paragraph(_pdf_markup(row), styles["meta"]))

    if template == "modern":
        story.extend(
            [
                HRFlowable(
                    width="100%", thickness=1.2, color=colors.HexColor("#" + config["accent"])
                ),
                Spacer(1, 5),
            ]
        )

    for line in resume_lines_for_export(record, section_order, hidden_sections):

        if not line:
            continue

        if _is_heading(line):
            story.append(Paragraph(_pdf_markup(line.rstrip(":")), styles["heading"]))

            story.append(
                HRFlowable(
                    width="100%", thickness=0.45, color=colors.HexColor("#B8C4CE"), spaceAfter=3
                )
            )

        else:
            prefix = "- " if line.startswith(("- ", "• ")) else ""

            content = line[2:].strip() if prefix else line

            story.append(Paragraph(prefix + _pdf_markup(content), styles["body"]))

    if include_sources:
        story.extend(
            [
                PageBreak(),
                Paragraph("ApplyEase Evidence Appendix", styles["heading"]),
                Paragraph(
                    "This appendix is for review and provenance; remove it before submitting the resume.",
                    styles["source"],
                ),
            ]
        )

        for source in (record.content or {}).get("sources", []):

            if not isinstance(source, dict):
                continue

            label = f"Experience #{source.get('experience_id')}: {source.get('experience_title', 'Confirmed experience')}"

            story.append(
                Paragraph(
                    "- <b>"
                    + _pdf_markup(label)
                    + "</b><br/>"
                    + _pdf_markup(str(source.get("text", ""))),
                    styles["source"],
                )
            )
    doc.build(story)
    return output.getvalue()
